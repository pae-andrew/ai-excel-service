"""Runnable checks for sandbox + multi-sheet I/O + output files. Run: python test_sandbox.py"""
import pandas as pd
import sandbox
import files

sheets = {"Data": pd.DataFrame({"a": [1, 1, 2], "email": ["x@y.z", "", "p@q.r"]})}

# 1. Inspection: print works, data unchanged.
out, _outs, stdout, err = sandbox.run(sheets, "print(len(df))")
assert err is None, err
assert stdout.strip() == "3"
assert len(out["Data"]) == 3

# 2. Mutation via df persists back to its sheet.
out, _outs, _, err = sandbox.run(sheets, "df = df[df['email'] != '']")
assert err is None, err
assert len(out["Data"]) == 2, out

# 3. Adding a new sheet via the sheets dict.
out, _outs, _, err = sandbox.run(sheets, "sheets['Counts'] = df.groupby('a').size().reset_index(name='n')")
assert err is None, err
assert "Counts" in out, list(out)

# 4. Blocked import -> reported, not executed.
_, _outs, _, err = sandbox.run(sheets, "import os; os.system('echo pwned')")
assert err and "blocked" in err, f"os should be blocked, got: {err!r}"

# 5. open() is read-only: writing a file is blocked.
_, _outs, _, err = sandbox.run(sheets, "open('/tmp/x', 'w')")
assert err and ("blocked" in err or "PermissionError" in err), f"write-open should be blocked, got: {err!r}"

# 6. Multi-sheet xlsx round-trip.
two = {"S1": pd.DataFrame({"x": [1, 2]}), "S2": pd.DataFrame({"y": [3]})}
back = files.read_to_sheets(files.sheets_to_xlsx_bytes(two), "r.xlsx")
assert set(back) == {"S1", "S2"}, list(back)
assert len(back["S1"]) == 2 and len(back["S2"]) == 1

# 7. save_result -> CSV output.
_, outs, _, err = sandbox.run(sheets, "save_result('out.csv', df.to_csv(index=False))")
assert err is None, err
assert outs and outs[-1][0] == "out.csv" and b"email" in outs[-1][1], outs

# 8. save_result -> PDF via matplotlib (Cyrillic-capable fonts).
code = (
    "import matplotlib; matplotlib.use('Agg')\n"
    "import matplotlib.pyplot as plt, io\n"
    "fig, ax = plt.subplots(); ax.axis('off')\n"
    "ax.table(cellText=df.values, colLabels=df.columns, loc='center')\n"
    "buf = io.BytesIO(); fig.savefig(buf, format='pdf'); save_result('report.pdf', buf)"
)
_, outs, _, err = sandbox.run(sheets, code)
assert err is None, err
assert outs and outs[-1][0] == "report.pdf" and outs[-1][1][:4] == b"%PDF", (err, outs and outs[-1][0])

# 9. save_result -> DOCX via python-docx.
code = (
    "import docx, io\n"
    "d = docx.Document(); d.add_heading('Отчёт', 0); d.add_paragraph('Привет')\n"
    "buf = io.BytesIO(); d.save(buf); save_result('doc.docx', buf)"
)
_, outs, _, err = sandbox.run(sheets, code)
assert err is None, err
assert outs and outs[-1][0] == "doc.docx" and outs[-1][1][:2] == b"PK", (err, outs and outs[-1][0])

# 10. Embed an uploaded image into a PDF via the `image` bytes (not redraw).
import struct as _st, zlib as _zl
def _png(w, h, rgb):
    raw = b"".join(b"\x00" + bytes(rgb) * w for _ in range(h))
    def ch(t, d):
        c = t + d; return _st.pack(">I", len(d)) + c + _st.pack(">I", _zl.crc32(c) & 0xffffffff)
    return (b"\x89PNG\r\n\x1a\n" + ch(b"IHDR", _st.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0))
            + ch(b"IDAT", _zl.compress(raw)) + ch(b"IEND", b""))
code = ("from PIL import Image\nimport io\n"
        "buf = io.BytesIO(); Image.open(io.BytesIO(image)).convert('RGB').save(buf, 'PDF')\n"
        "save_result('shot.pdf', buf)")
_, outs, _, err = sandbox.run({}, code, images=[_png(6, 6, (10, 200, 90))])
assert err is None, err
assert outs and outs[-1][0] == "shot.pdf" and outs[-1][1][:4] == b"%PDF", (err, outs and outs[-1][0])

print("OK: sandbox + multi-sheet + output-files (csv/pdf/docx) + image-embed self-check passed")
